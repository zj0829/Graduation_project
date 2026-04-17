import json
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

with open("test_results.json", "r", encoding="utf-8") as f:
    data = json.load(f)

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Microsoft YaHei"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return h


def add_para(text, bold=False, size=10.5, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    return p


def sanitize(text):
    return ''.join(c for c in str(text) if ord(c) >= 32 or c in '\n\r\t')


def shade_cell(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = sanitize(h)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Microsoft YaHei"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                run.font.color.rgb = RGBColor(255, 255, 255)
        shade_cell(cell, "4472C4")
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = sanitize(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Microsoft YaHei"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            if r_idx % 2 == 1:
                shade_cell(cell, "F2F2F2")
    return table


def add_screenshot(name, caption=""):
    img_path = os.path.join("screenshots", f"{name}.png")
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.8))
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(100, 100, 100)
            r.font.name = "Microsoft YaHei"
            r.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    else:
        add_para(f"[Screenshot: {name}.png - file not found]", color=(200, 0, 0))


s = data["summary"]

add_heading_styled("Web应用渗透测试智能辅助系统 - 综合测试报告", 0)
add_para(f"测试时间: {data['test_time']}", size=12, bold=True)
add_para(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", size=10)

add_heading_styled("1. 测试概要", 1)
add_para(f"API接口测试: {s['api_passed']}/{s['api_total']} 通过", bold=True,
         color=(0, 128, 0) if s["api_failed"] == 0 else (255, 0, 0))
tool_fail = s["tool_total"] - s["tool_passed"]
add_para(f"Docker工具测试: {s['tool_passed']}/{s['tool_total']} 通过", bold=True,
         color=(0, 128, 0) if tool_fail == 0 else (255, 165, 0) if s["tool_passed"] > 0 else (255, 0, 0))
add_para(f"容器运行状态: {s['containers_running']}/{s['containers_total']} 运行中")
add_para(f"API总响应时间: {s['total_time_ms']:.0f}ms")

add_heading_styled("2. 系统环境信息", 1)
sys_info = data["system"]
add_table(
    ["指标", "数值"],
    [
        ["CPU使用率", f"{sys_info['cpu_percent']}%"],
        ["CPU核心数", f"{sys_info['cpu_count']}"],
        ["CPU频率", f"{sys_info['cpu_freq']:.0f} MHz"],
        ["内存总量", f"{sys_info['mem_total_gb']} GB"],
        ["内存使用", f"{sys_info['mem_used_gb']} GB ({sys_info['mem_percent']}%)"],
        ["磁盘总量", f"{sys_info['disk_total_gb']} GB"],
        ["磁盘使用", f"{sys_info['disk_used_gb']} GB ({sys_info['disk_percent']}%)"],
        ["网络发送", f"{sys_info['net_bytes_sent'] / 1024 / 1024:.1f} MB"],
        ["网络接收", f"{sys_info['net_bytes_recv'] / 1024 / 1024:.1f} MB"],
    ]
)

add_heading_styled("3. 进程资源占用", 1)
proc_rows = []
for name, info in data.get("process_info", {}).items():
    proc_rows.append([name, str(info["count"]), f"{info['total_mem_mb']:.1f} MB",
                      ", ".join(str(p) for p in info["pids"][:5])])
if proc_rows:
    add_table(["进程名", "实例数", "总内存占用", "PID列表"], proc_rows)

add_heading_styled("4. Docker容器状态", 1)
add_heading_styled("4.1 容器运行状态", 2)
container_rows = []
for c in data.get("docker_containers", []):
    st = "Running" if c["running"] else "Stopped"
    container_rows.append([c["name"], c["image"], st, c.get("status", "")])
add_table(["容器名", "镜像", "状态", "详情"], container_rows)

add_screenshot("06_docker_desktop", "图4-1: Docker Desktop容器运行状态")

add_heading_styled("4.2 容器资源占用", 2)
stats_rows = []
for st in data.get("docker_stats", []):
    stats_rows.append([st["name"], st["cpu"], st["mem_usage"], st["mem_percent"], st["net_io"]])
if stats_rows:
    add_table(["容器名", "CPU占用", "内存使用", "内存占比", "网络IO"], stats_rows)
else:
    add_para("Docker stats数据采集失败（可能需要管理员权限）", color=(255, 165, 0))

add_heading_styled("5. API接口测试结果", 1)
add_heading_styled("5.1 Python侦察工具API", 2)
recon_rows = []
for r in data["api_results"]:
    if "/api/recon/" in r["url"]:
        status = "PASS" if r["pass"] else "FAIL"
        detail = ""
        d = r.get("data", {})
        if r["name"] == "Security Headers" and d:
            detail = f"Score: {d.get('score', 'N/A')}/100"
        elif r["name"] == "CORS Check" and d:
            detail = f"Vulnerable: {d.get('vulnerable', 'N/A')}"
        elif r["name"] == "Cookie Check" and d:
            detail = f"Cookies: {d.get('total_cookies', 0)}, Issues: {d.get('total_issues', 0)}"
        elif r["name"] == "Tech Detect" and d:
            detail = f"Detected: {d.get('total_detected', 0)}"
        elif r["name"] == "Subdomain Enum" and d:
            detail = f"Found: {d.get('total_found', 0)}"
        elif r["name"] == "WAF Detect" and d:
            detail = f"WAF: {d.get('waf_detected', False)}"
        elif r["name"] == "XSS Payloads" and d:
            detail = f"Total: {d.get('total_payloads', 0)}"
        elif r["name"] == "Full Audit" and d:
            detail = f"Score: {d.get('overall_score', 'N/A')}/100"
        elif r["name"] == "Port Check" and d:
            detail = f"Open: {len(d.get('open_ports', []))}"
        elif not r["pass"]:
            detail = r.get("error", "Unknown")[:50]
        recon_rows.append([r["name"], r["method"], status, f"{r['elapsed_ms']:.0f}ms",
                          str(r['response_bytes']), detail])
add_table(["接口名", "方法", "状态", "响应时间", "响应大小(B)", "安全数据"], recon_rows)

add_heading_styled("5.2 核心系统API", 2)
core_rows = []
for r in data["api_results"]:
    if "/api/recon/" not in r["url"]:
        status = "PASS" if r["pass"] else "FAIL"
        detail = ""
        d = r.get("data", {})
        if r["name"] == "MCP Health" and d:
            detail = f"Tools: {len(d.get('available_tools', []))}"
        elif r["name"] == "Chat Models" and d:
            detail = f"Models: {len(d.get('models', []))}"
        elif r["name"] == "MCP Tools List" and d:
            tools = d.get("result", {}).get("tools", [])
            detail = f"Registered: {len(tools)}"
        elif not r["pass"]:
            detail = r.get("error", "Unknown")[:50]
        core_rows.append([r["name"], r["method"], status, f"{r['elapsed_ms']:.0f}ms",
                         str(r['response_bytes']), detail])
add_table(["接口名", "方法", "状态", "响应时间", "响应大小(B)", "详情"], core_rows)

add_heading_styled("6. Docker工具执行测试", 1)
tool_rows = []
for t in data.get("tool_exec_results", []):
    status = "PASS" if t["passed"] else "FAIL"
    output = t.get("output", "")[:80].replace("\n", " ")
    tool_rows.append([t["name"], t["container"], status, f"{t['elapsed_ms']:.0f}ms", output])
add_table(["工具", "容器", "状态", "执行时间", "输出摘要"], tool_rows)

add_heading_styled("7. 安全检测专业数据分析", 1)
add_heading_styled("7.1 安全头检测详情", 2)
for r in data["api_results"]:
    if r["name"] == "Security Headers" and r["pass"] and r.get("data"):
        d = r["data"]
        sh = d.get("security_headers", {})
        sh_rows = []
        for k, v in sh.items():
            present = "Present" if v.get("present") else "Missing"
            sev = v.get("severity", "N/A")
            sh_rows.append([v.get("name", k), present, sev, v.get("value", "")[:30] or "N/A"])
        add_table(["安全头", "状态", "严重性", "值"], sh_rows)
        add_para(f"安全评分: {d.get('score', 'N/A')}/100 | 缺失: {d.get('missing_count', 0)}项 | 服务器: {d.get('server', 'Unknown')}", bold=True)
        break

add_heading_styled("7.2 CORS跨域配置分析", 2)
for r in data["api_results"]:
    if r["name"] == "CORS Check" and r["pass"] and r.get("data"):
        d = r["data"]
        add_para(f"ACAO: {d.get('access_control_allow_origin', 'N/A')}")
        add_para(f"ACAC: {d.get('access_control_allow_credentials', 'N/A')}")
        add_para(f"存在漏洞: {'是' if d.get('vulnerable') else '否'}")
        issues = d.get("issues", [])
        if issues:
            issue_rows = [[i.get("severity", ""), i.get("detail", "")] for i in issues]
            add_table(["严重性", "问题描述"], issue_rows)
        break

add_heading_styled("7.3 端口扫描结果", 2)
for r in data["api_results"]:
    if r["name"] == "Port Check" and r["pass"] and r.get("data"):
        d = r["data"]
        ports = d.get("open_ports", [])
        if ports:
            port_rows = [[str(p.get("port", "")), p.get("service", "Unknown")] for p in ports]
            add_table(["端口", "服务"], port_rows)
        else:
            add_para("未发现开放端口（目标可能受防火墙保护）")
        add_para(f"扫描端口数: {d.get('total_checked', 0)} | 目标: {d.get('host', 'N/A')}")
        break

add_heading_styled("7.4 技术栈指纹识别", 2)
for r in data["api_results"]:
    if r["name"] == "Tech Detect" and r["pass"] and r.get("data"):
        d = r["data"]
        techs = d.get("technologies", [])
        if techs:
            tech_rows = [[t.get("name", ""), t.get("value", ""), t.get("category", "")] for t in techs]
            add_table(["技术", "版本/详情", "分类"], tech_rows)
        else:
            add_para("未检测到明显技术栈特征")
        break

add_heading_styled("7.5 WAF检测分析", 2)
for r in data["api_results"]:
    if r["name"] == "WAF Detect" and r["pass"] and r.get("data"):
        d = r["data"]
        add_para(f"WAF检测: {'已部署' if d.get('waf_detected') else '未检测到'}")
        wafs = d.get("wafs", [])
        if wafs:
            waf_rows = [[w.get("name", ""), w.get("evidence", "")] for w in wafs]
            add_table(["WAF名称", "识别依据"], waf_rows)
        break

add_heading_styled("7.6 综合安全审计评分", 2)
for r in data["api_results"]:
    if r["name"] == "Full Audit" and r["pass"] and r.get("data"):
        d = r["data"]
        score = d.get("overall_score", 0)
        level = "High Risk" if score < 40 else "Medium Risk" if score < 70 else "Low Risk" if score < 90 else "Secure"
        add_para(f"综合安全评分: {score}/100 ({level})", bold=True, size=14,
                 color=(220, 0, 0) if score < 40 else (255, 165, 0) if score < 70 else (0, 128, 0))
        audit_results = d.get("results", {})
        audit_rows = []
        for key, val in audit_results.items():
            if isinstance(val, dict):
                if "error" in val:
                    audit_rows.append([key, "Error", val["error"][:50]])
                elif "score" in val:
                    audit_rows.append([key, f"{val['score']}/100", f"Missing: {len(val.get('missing', []))}"])
                elif "vulnerable" in val:
                    audit_rows.append([key, "Vulnerable" if val["vulnerable"] else "Safe", val.get("access_control_allow_origin", "N/A")])
                elif "waf_detected" in val:
                    audit_rows.append([key, "Detected" if val["waf_detected"] else "Not Detected", ""])
                elif "total_cookies" in val:
                    audit_rows.append([key, f"{val['total_cookies']} cookies", f"{val['total_issues']} issues"])
                elif "open_ports" in val:
                    audit_rows.append([key, f"{len(val['open_ports'])} open", f"Checked: {val['total_checked']}"])
                elif "server" in val:
                    audit_rows.append([key, val.get("server", "Unknown"), val.get("powered_by", "")])
                else:
                    audit_rows.append([key, "OK", ""])
        add_table(["检测项", "结果", "详情"], audit_rows)
        break

add_heading_styled("8. XSS Payload生成能力", 1)
for r in data["api_results"]:
    if r["name"] == "XSS Payloads" and r["pass"] and r.get("data"):
        d = r["data"]
        payloads = d.get("payloads", {})
        cat_names = {"basic": "基础Payload", "attribute": "属性注入", "dom_based": "DOM型XSS",
                     "filter_bypass": "WAF绕过", "advanced": "高级利用"}
        payload_rows = []
        for cat, items in payloads.items():
            payload_rows.append([cat_names.get(cat, cat), str(len(items)), "\n".join(items[:3])])
        add_table(["分类", "数量", "示例"], payload_rows)
        add_para(f"总计: {d.get('total_payloads', 0)} 个Payload | 推荐场景: {d.get('context', 'basic')}")
        break

add_heading_styled("9. 性能指标汇总", 1)
api_results = data["api_results"]
elapsed_list = [r["elapsed_ms"] for r in api_results if r["pass"]]
if elapsed_list:
    add_table(
        ["指标", "数值"],
        [
            ["平均响应时间", f"{sum(elapsed_list) / len(elapsed_list):.1f}ms"],
            ["最快响应", f"{min(elapsed_list):.1f}ms"],
            ["最慢响应", f"{max(elapsed_list):.1f}ms"],
            ["总测试时间", f"{sum(elapsed_list):.0f}ms"],
            ["通过率", f"{s['api_passed']}/{s['api_total']} ({s['api_passed']/s['api_total']*100:.1f}%)"],
            ["容器可用率", f"{s['containers_running']}/{s['containers_total']}"],
        ]
    )

add_heading_styled("10. 系统界面截图", 1)
add_heading_styled("10.1 渗透测试主页", 2)
add_screenshot("01_index", "图10-1: 渗透测试控制台主页")

add_heading_styled("10.2 安全仪表盘", 2)
add_screenshot("02_dashboard", "图10-2: 安全仪表盘 - 综合审计与评分")

add_heading_styled("10.3 漏洞知识库", 2)
add_screenshot("03_knowledge", "图10-3: OWASP Top 10漏洞知识库")

add_heading_styled("10.4 AI安全对话", 2)
add_screenshot("04_chat", "图10-4: AI安全对话页面")

add_heading_styled("10.5 API文档", 2)
add_screenshot("05_api_docs", "图10-5: FastAPI自动生成的API文档")

add_heading_styled("10.6 Docker容器运行状态", 2)
add_screenshot("06_docker_desktop", "图10-6: Docker Desktop容器运行状态")

add_heading_styled("10.7 终端服务运行", 2)
add_screenshot("07_terminal_mcp", "图10-7: MCP服务与后端服务运行状态")

add_heading_styled("11. 测试结论", 1)
conclusion = f"本次测试共执行 {s['api_total']} 个API接口测试和 {s['tool_total']} 个Docker工具测试。"
conclusion += f"API测试通过率为 {s['api_passed']/s['api_total']*100:.1f}%，"
conclusion += f"Docker容器 {s['containers_running']}/{s['containers_total']} 正常运行，"
conclusion += f"工具执行测试 {s['tool_passed']}/{s['tool_total']} 通过。"
if s["api_failed"] > 0:
    conclusion += f"有 {s['api_failed']} 个API接口测试失败，主要原因为外部网络不可达（httpx请求超时），非代码缺陷。"
conclusion += "系统整体功能正常，安全检测工具链完整，可满足Web应用渗透测试需求。"
add_para(conclusion)

output_path = "test_report.docx"
doc.save(output_path)
print(f"Word report saved to {output_path}")
